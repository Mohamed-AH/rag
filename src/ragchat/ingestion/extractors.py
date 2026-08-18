"""Extract structured :class:`Section` objects from uploaded files.

Dispatches on file extension: markdown keeps its heading structure; plain text, PDF, and
Word documents are text-extracted and then size-chunked. All formats funnel into the same
``Section`` pipeline the rest of the app already understands.
"""

from __future__ import annotations

import io
from pathlib import PurePosixPath

from ragchat.errors import EmptyDocumentError, UnsupportedFileTypeError
from ragchat.ingestion.chunker import chunk_text
from ragchat.ingestion.parser import Section, parse_markdown

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
    {".md", ".markdown", ".txt", ".text", ".pdf", ".docx"}
)


def _pdf_to_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def pdf_to_png_pages(data: bytes, *, max_pages: int = 10, dpi: int = 150) -> list[bytes]:
    """Rasterize a PDF's pages to PNG images (one per page, up to ``max_pages``).

    Scanned PDFs go to the multimodal model as images. Gemini accepts an inline PDF, but
    Mistral/Groq/OpenAI-compatible vision models only accept real image formats — so we
    render to PNG, which every provider accepts. Uses pypdfium2 (BSD-licensed, wheel-bundled
    binaries — no AGPL, no system poppler) + Pillow, both imported lazily.
    """
    import pypdfium2 as pdfium
    from PIL import Image  # noqa: F401  (ensures Pillow is present for to_pil)

    scale = dpi / 72.0
    pngs: list[bytes] = []
    pdf = pdfium.PdfDocument(data)
    try:
        for index in range(min(len(pdf), max_pages)):
            bitmap = pdf[index].render(scale=scale)
            image = bitmap.to_pil().convert("RGB")
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            pngs.append(buffer.getvalue())
    finally:
        pdf.close()
    return pngs


def pack_images(images: list[bytes], max_images: int, *, max_height: int = 4000) -> list[bytes]:
    """Combine ``images`` into at most ``max_images`` PNGs by vertically stacking pages.

    Some vision providers cap the number of images per request (e.g. Groq's Qwen-VL allows 3).
    Rather than drop pages — which would drop whole documents from a packet — we pack pages
    into ``max_images`` composite images (contiguous groups, each stacked top-to-bottom), so
    every page's content still reaches the model. Composites taller than ``max_height`` are
    downscaled to stay within provider pixel limits. A no-op when already within the cap.
    """
    if max_images <= 0 or len(images) <= max_images:
        return images

    from PIL import Image

    per_group = -(-len(images) // max_images)  # ceil division → at most max_images groups
    groups = [images[i : i + per_group] for i in range(0, len(images), per_group)]
    packed: list[bytes] = []
    for group in groups:
        tiles = [Image.open(io.BytesIO(b)).convert("RGB") for b in group]
        width = max(tile.width for tile in tiles)
        resized = [
            tile
            if tile.width == width
            else tile.resize((width, round(tile.height * width / tile.width)))
            for tile in tiles
        ]
        total_height = sum(tile.height for tile in resized)
        canvas = Image.new("RGB", (width, total_height), "white")
        offset = 0
        for tile in resized:
            canvas.paste(tile, (0, offset))
            offset += tile.height
        if canvas.height > max_height:
            scale = max_height / canvas.height
            canvas = canvas.resize((max(1, round(width * scale)), max_height))
        buffer = io.BytesIO()
        canvas.save(buffer, format="PNG")
        packed.append(buffer.getvalue())
    return packed


def _docx_to_text(data: bytes) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(data))
    return "\n\n".join(p.text for p in document.paragraphs if p.text.strip())


def extract_sections(
    filename: str,
    data: bytes,
    *,
    chunk_max_chars: int = 1200,
    chunk_overlap: int = 150,
) -> list[Section]:
    """Parse ``data`` (an uploaded file's bytes) into sections based on ``filename``.

    Raises :class:`UnsupportedFileTypeError` for unknown extensions and
    :class:`EmptyDocumentError` when no text can be extracted.
    """
    suffix = PurePosixPath(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix or filename}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    label = PurePosixPath(filename).name

    if suffix in {".md", ".markdown"}:
        sections = parse_markdown(data.decode("utf-8", errors="replace"))
    else:
        if suffix in {".txt", ".text"}:
            text = data.decode("utf-8", errors="replace")
        elif suffix == ".pdf":
            text = _pdf_to_text(data)
        else:  # .docx
            text = _docx_to_text(data)
        sections = chunk_text(text, source=label, max_chars=chunk_max_chars, overlap=chunk_overlap)

    if not sections:
        raise EmptyDocumentError(f"No extractable text found in '{label}'.")
    return sections
