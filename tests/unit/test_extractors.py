"""Tests for multi-format file extraction."""

from __future__ import annotations

import io

import pytest

from ragchat.errors import EmptyDocumentError, UnsupportedFileTypeError
from ragchat.ingestion import extractors
from ragchat.ingestion.extractors import extract_sections


def test_markdown_uses_headings() -> None:
    data = b"# Heading One\nbody a\n\n# Heading Two\nbody b\n"
    sections = extract_sections("notes.md", data)
    assert [s.title for s in sections] == ["Heading One", "Heading Two"]


def test_plain_text_is_chunked_with_source_title() -> None:
    sections = extract_sections("notes.txt", b"Some plain text content.")
    assert len(sections) == 1
    assert sections[0].title == "notes.txt (part 1)"


def test_docx_is_extracted() -> None:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("First paragraph of the report.")
    doc.add_paragraph("Second paragraph with more detail.")
    buf = io.BytesIO()
    doc.save(buf)

    sections = extract_sections("report.docx", buf.getvalue())
    assert len(sections) >= 1
    assert "First paragraph" in sections[0].content


def test_pdf_dispatch(monkeypatch: pytest.MonkeyPatch) -> None:
    # Extracting text from a real PDF is pypdf's job; here we verify the .pdf branch
    # routes through the chunker with a stubbed extractor.
    monkeypatch.setattr(extractors, "_pdf_to_text", lambda data: "Extracted PDF text about VPCs.")
    sections = extract_sections("doc.pdf", b"%PDF-1.4 fake bytes")
    assert len(sections) == 1
    assert "VPC" in sections[0].content
    assert sections[0].title == "doc.pdf (part 1)"


def test_unsupported_extension_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        extract_sections("malware.exe", b"MZ...")


def test_empty_document_raises() -> None:
    with pytest.raises(EmptyDocumentError):
        extract_sections("blank.txt", b"    \n\n   ")
